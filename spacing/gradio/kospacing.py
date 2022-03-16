import gradio as gr
from pykospacing import Spacing
from docx import Document

# kospacing
def kospacing(file):
    spacing = Spacing()
    
    document = Document(file)
    new_doc = Document()
    para = new_doc.add_paragraph()
    
    for paragraph in document.paragraphs:
        text = paragraph.text
            
        try:
            space_sent = spacing(text)
            para.add_run(space_sent+"\n")
        except:
            para.add_run(text+"\n")
    
    #text = docx2txt.process(file)
    #space_sent = spacing(text)
    #para.add_run(text)
    
    new_doc.save("ko_result.docx")
    return "ko_result.docx"

ko_iface = gr.Interface(fn=kospacing, inputs="file", outputs="file", title="Kospacing")
ko_iface.launch(server_name="0.0.0.0", server_port=8000, share=False)